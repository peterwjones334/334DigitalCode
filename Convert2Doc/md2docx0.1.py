# md2docx0.1.py
# version 0.1
# 03-november 2023

"""
check regex
learn python-docx
"""

import re
from docx import Document

# Define simple patterns for Markdown
HEADING_PATTERN = re.compile(r'^(#{1,6})\s*(.*)')
LIST_PATTERN = re.compile(r'^(\*|\d+\.)\s*(.*)')
ITALIC_BOLD_PATTERN = re.compile(r'(\*\*\*|___)(.*?)\1')
BOLD_PATTERN = re.compile(r'(\*\*|__)(.*?)\1')
ITALIC_PATTERN = re.compile(r'(\*|_)(.*?)\1')
INLINE_CODE_PATTERN = re.compile(r'(`)(.*?)\1')

class MarkdownConverter:
    def __init__(self, md_text):
        self.md_text = md_text
        self.doc = Document()

    def apply_text_styles(self, paragraph, text):
        # This simplified version just replaces markdown symbols with nothing
        text = ITALIC_BOLD_PATTERN.sub(r'\2', text)
        text = BOLD_PATTERN.sub(r'\2', text)
        text = ITALIC_PATTERN.sub(r'\2', text)
        text = INLINE_CODE_PATTERN.sub(r'\2', text)
        paragraph.add_run(text)

    def convert(self):
        for line in self.md_text.split('\n'):
            if line.strip() == '':
                continue

            heading_match = HEADING_PATTERN.match(line)
            list_match = LIST_PATTERN.match(line)

            if heading_match:
                level = len(heading_match.group(1))
                self.doc.add_heading(heading_match.group(2), level=level - 1)
            elif list_match:
                self.doc.add_paragraph(list_match.group(2), style='ListBullet' if list_match.group(1).startswith('*') else 'ListNumber')
            else:
                para = self.doc.add_paragraph()
                self.apply_text_styles(para, line)

        return self.doc

# Example usage - write or paste md file here between quotes

md_text = """
# Heading 1
## Heading 2
* List item 1
* List item 2
1. Numbered item 1
2. Numbered item 2
**Bold text** and _italic text_ with `inline code`.
"""
converter = MarkdownConverter(md_text)
doc = converter.convert()
doc.save('simplified_output.docx')
