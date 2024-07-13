## md2docx1.py
## version 1
## 02-nov-2023

"""

The code provides a script that reads a Markdown file, parses various elements (headings, bullets, hyperlinks, etc.), and converts it into a Word document using the python-docx library.
It defines functions to handle individual Markdown elements and uses regular expressions to detect these elements in the text.

Here's an explanation of some key parts of the script:

- parse_markdown(md_text): This is the main function that takes Markdown text as input and converts it to a Word document.
- Regular expressions are used to identify Markdown patterns (headings, bullets, etc.).
- Markdown elements are converted to Word document elements through functions like add_heading, add_paragraph, add_bullet, add_blockquote, etc.
- The add_hyperlink function uses low-level XML manipulation to add hyperlinks because python-docx may not support them directly.

To use this script, you need to have a Markdown file saved as example.md in the same directory as your script.
When you run the script, it should generate a Word document named example.docx with the content formatted according to the Markdown syntax.
The script covers basic Markdown elements. However, it does not support more complex elements like tables or images out of the box.
If you encounter any errors or need to handle additional Markdown features, you would need to expand the script accordingly.

Before running the script, ensure you have the python-docx module installed in your Python environment:
>pip install python-docx

Also, if your Word version doesn't have a style named 'ListBullet' or 'Intense Quote', you may need to define these styles in the document or replace them with existing ones.

"""

import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml import ns

# Define how each markdown element should be handled
def add_heading(document, line, level):
    document.add_heading(line, level=level)

def add_paragraph(document, line):
    document.add_paragraph(line)

def add_bullet(document, line):
    document.add_paragraph(line, style='ListBullet')

# Parse the markdown text line by line
def parse_markdown(md_text):
    document = Document()

    # Regular expressions for markdown patterns
    heading_pattern = re.compile(r'^(#{1,6})\s*(.*)')
    bullet_pattern = re.compile(r'^\*\s+(.*)')
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    italic_pattern = re.compile(r'\*(.*?)\*')
    blockquote_pattern = re.compile(r'^>\s*(.*)')
    hyperlink_pattern = re.compile(r'\[(.*?)\]\((.*?)\)')

    for line in md_text.split('\n'):
        # Check for blockquotes
        blockquote_match = blockquote_pattern.match(line)
        if blockquote_match:
            add_blockquote(document, blockquote_match.group(1))
        # Check for headings
        elif heading_match := heading_pattern.match(line):
            level = len(heading_match.group(1))  # Number of # determines level
            text = heading_match.group(2).strip()
            document.add_heading(text, level)
        # Check for bullets
        elif bullet_match := bullet_pattern.match(line):
            text = bullet_pattern.sub(r'\1', line)
            document.add_paragraph(text, style='ListBullet')
        # Check for hyperlinks
        elif hyperlink_match := hyperlink_pattern.search(line):
            paragraph = document.add_paragraph()
            start = 0
            for match in hyperlink_pattern.finditer(line):
                # Add the text before the hyperlink
                paragraph.add_run(line[start:match.start()])
                # Add the hyperlink
                add_hyperlink(paragraph, match.group(2), match.group(1))
                # Update the start to the end of the current match
                start = match.end()
            # Add any remaining text after the last hyperlink
            paragraph.add_run(line[start:])
        # Ordinary paragraph
        else:
            paragraph = document.add_paragraph()
            apply_text_styles(paragraph, line)

    return document

# Helper function to apply bold and italic styles
def apply_text_styles(paragraph, line):
    # Patterns for bold and italic
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    italic_pattern = re.compile(r'(?<!\*)\*(?!\*)(.*?)\*(?!\*)')
    code_inline_pattern = re.compile(r'`(.*?)`')

    # Find all style patterns
    bold_matches = list(re.finditer(bold_pattern, line))
    italic_matches = list(re.finditer(italic_pattern, line))

    # Sort all matches by start position
    all_matches = sorted(bold_matches + italic_matches, key=lambda match: match.start())

    # Initial cursor position
    cursor = 0

    # Apply styles
    for match in all_matches:
        style_start, style_end = match.span()
        # Add the text before the styled part
        if cursor < style_start:
            paragraph.add_run(line[cursor:style_start])
        # Add the styled part
        styled_run = paragraph.add_run(match.group(1))
        if match in bold_matches:
            styled_run.bold = True
        elif match in italic_matches:
            styled_run.italic = True
        # Move the cursor past this styled part
        cursor = style_end

    # Add any remaining text after the last style pattern
    if cursor < len(line):
        paragraph.add_run(line[cursor:])

# Function to add a hyperlink to a paragraph
def add_hyperlink(paragraph, url, text):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'hyperlink', is_external=True)

    # Create the w:hyperlink tag and set needed attribute
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(ns.qn('r:id'), r_id)

    # Create a w:r element for run properties
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')  # Run properties

    # Set the style for hyperlinks
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(ns.qn('w:val'), 'Hyperlink')  # 'Hyperlink' is a default style in Word for hyperlinks

    # Join all the xml elements together and add the required text to the w:r element
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink

# Function to handle blockquotes
def add_blockquote(document, line):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles['Intense Quote']
    paragraph.add_run(line)

# Load a markdown file and convert it to a word document
def convert_md_to_docx(md_filepath, docx_filepath):
    with open(md_filepath, 'r', encoding='utf-8') as file:
        md_text = file.read()

    document = parse_markdown(md_text)
    document.save(docx_filepath)
    print(f"Converted {md_filepath} to {docx_filepath} successfully.")

# Example usage
md_file = 'example.md'  # Your markdown file
docx_file = 'example.docx'  # The Word file you want to create
convert_md_to_docx(md_file, docx_file)