# md2docx2.py
# version 2.0
# 02-november 2023

"""
This code is a Python script designed to parse a Markdown file and convert it into a .docx file.
 It utilizes the python-docx library to create and manipulate Word documents. 
 The script identifies various Markdown elements like headings, bullets, blockquotes, hyperlinks, bold and italic text, inline code, and code blocks.

Each function is designed to handle a specific Markdown pattern:

- parse_markdown(md_text): This is the main function that goes through each line of Markdown text, determines what elements are present using other functions, and builds the Word document.
- check_heading(line, document): Checks for Markdown headings and adds them to the Word document with appropriate levels.
- check_bullet(line, document): Looks for bullet points and adds them to the document with the 'ListBullet' style.
= check_blockquote(line, document): Identifies blockquotes and applies the 'Intense Quote' style in the document.
- check_hyperlink(line, document): Detects hyperlinks in the text and properly formats them with a function that creates hyperlink XML elements.
- apply_text_styles(paragraph, line): Applies inline text styles such as bold, italic, and inline code based on Markdown syntax.
- add_code_block_to_document(document, code_block_content): Adds a code block to the document, formatting it with a monospace font and ensuring it keeps its distinct appearance.
- add_hyperlink(paragraph, url, text): A helper function to add a hyperlink to a paragraph with proper Word XML structure.
- convert_md_to_docx(md_filepath, docx_filepath): Opens a Markdown file, parses it, and saves the result to a .docx file.

The example.md file will be read, parsed, and the output will be saved to example.docx when you run this script.

Make sure you have the python-docx library installed in your Python environment to use this script. 
If not, you can install it using pip:
pip install python-docx

Remember that the script defines some patterns at the top that are used to identify Markdown elements using regular expressions.
If your Markdown files use a different syntax or have elements not covered by this script, you may need to adjust the patterns and functions accordingly.
"""

import re
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn  # This is used to correctly handle XML namespaces

# Global constants for patterns
HEADING_PATTERN = re.compile(r'^(#{1,6})\s*(.*)')
BULLET_PATTERN = re.compile(r'^\*\s+(.*)')
BLOCKQUOTE_PATTERN = re.compile(r'^>\s*(.*)')
HYPERLINK_PATTERN = re.compile(r'\[(.*?)\]\((.*?)\)')

# Function to handle different markdown elements
def parse_markdown(md_text):
    document = Document()
    inside_code_block = False
    code_block_content = []

    for line in md_text.split('\n'):
        line_stripped = line.strip()
        if line_stripped.startswith('```'):  # Check for code block markers
            if inside_code_block:
                # End of code block
                inside_code_block = False
                add_code_block_to_document(document, code_block_content)
                code_block_content = []  # Clear the content after adding to the document
            else:
                inside_code_block = True
            continue  # Skip the ``` line

        if inside_code_block:
            code_block_content.append(line)
            continue

        # Refactored into separate functions for better readability and maintainability
        if check_heading(line, document):
            continue
        if check_bullet(line, document):
            continue
        if check_blockquote(line, document):
            continue
        if check_hyperlink(line, document):
            continue
        if line_stripped:  # Non-empty line
            paragraph = document.add_paragraph()
            apply_text_styles(paragraph, line)

    return document  # This should be indented inside the function to return the Document object

def check_heading(line, document):
    match = HEADING_PATTERN.match(line)
    if match:
        level = len(match.group(1))
        document.add_heading(match.group(2), level=level)
        return True
    return False

def check_bullet(line, document):
    match = BULLET_PATTERN.match(line)
    if match:
        paragraph = document.add_paragraph(style='ListBullet')
        apply_text_styles(paragraph, match.group(1))
        return True
    return False

def check_blockquote(line, document):
    match = BLOCKQUOTE_PATTERN.match(line)
    if match:
        paragraph = document.add_paragraph(style='Intense Quote')
        apply_text_styles(paragraph, match.group(1))
        return True
    return False

def check_hyperlink(line, document):
    matches = list(HYPERLINK_PATTERN.finditer(line))
    if matches:
        paragraph = document.add_paragraph()
        last_idx = 0
        for match in matches:
            # Add text before the hyperlink
            if last_idx < match.start():
                paragraph.add_run(line[last_idx:match.start()])
            # Add the hyperlink
            add_hyperlink(paragraph, match.group(2), match.group(1))
            last_idx = match.end()
        # Add any remaining text after the last hyperlink
        if last_idx < len(line):
            paragraph.add_run(line[last_idx:])
        return True
    return False

# Function to apply text styles like bold and italic
def apply_text_styles(paragraph, line):
    # Define patterns for bold and italic
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    italic_pattern = re.compile(r'(?<!\*)\*(?!\*)(.*?)\*(?!\*)')
    code_inline_pattern = re.compile(r'`(.*?)`')
    
    # Find all matches and sort them by start position
    bold_matches = list(bold_pattern.finditer(line))
    italic_matches = list(italic_pattern.finditer(line))
    code_inline_matches = list(code_inline_pattern.finditer(line))
    all_matches = sorted(bold_matches + italic_matches + code_inline_matches, key=lambda match: match.start())
    
    cursor = 0  # Cursor to keep track of where we are in the line
    for match in all_matches:
        start, end = match.span()
        # Add text before styled text
        if cursor < start:
            paragraph.add_run(line[cursor:start])
        # Style the text according to the markdown
        styled_text = match.group(1)
        run = paragraph.add_run(styled_text)
        if match in bold_matches:
            run.bold = True
        elif match in italic_matches:
            run.italic = True
        elif match in code_inline_matches:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        cursor = end  # Move cursor to the end of the styled text
    
    # Add any remaining text after the last style
    if cursor < len(line):
        paragraph.add_run(line[cursor:])

def add_code_block_to_document(document, code_block_content):
    # Create a new paragraph for the code block with a specific style
    paragraph = document.add_paragraph(style='No Spacing')
    # Add each line of the code block as a separate run
    for code_line in code_block_content:
        run = paragraph.add_run(code_line + '\n')
        run.font.name = 'Consolas'  # Set a monospace font for code blocks
        run.font.size = Pt(10)  # Set font size

    # Clear the content after it's added to the document
    code_block_content.clear()        

# Function to add a hyperlink to a paragraph
def add_hyperlink(paragraph, url, text):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)  

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a w:rPr (run properties) element
    rPr = OxmlElement('w:rPr')

    # Set the style for the run properties
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink

# Function to convert a markdown file to a docx file
def convert_md_to_docx(md_filepath, docx_filepath):
    with open(md_filepath, 'r', encoding='utf-8') as file:
        md_text = file.read()

    document = parse_markdown(md_text)
    document.save(docx_filepath)
    print(f"Converted {md_filepath} to {docx_filepath} successfully.")

# Example usage
md_file = 'example.md'  # Your markdown file
docx_file = 'example.docx'  # The Word file you want to create
convert_md_to_docx(md_file, docx_file)