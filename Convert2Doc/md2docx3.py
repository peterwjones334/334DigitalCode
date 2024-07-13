# md2docx3.py
# version 3.0
# 03-november 2023

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from urllib.request import urlopen
from io import BytesIO

# Define global constants for regex patterns
HEADING_PATTERN = re.compile(r'^(#{1,6})\s*(.*)')
LIST_PATTERN = re.compile(r'^(\*|\d+\.)\s+(.*)')
# Regular expression patterns for detecting Markdown lists
UNORDERED_LIST_PATTERN = re.compile(r'^(\s*)([*\-+])\s+(.*)')
ORDERED_LIST_PATTERN = re.compile(r'^(\s*)\d+\.\s+(.*)')
# BULLET_PATTERN = re.compile(r'^(\*|-|\+)\s+(.*)')
BULLET_PATTERN = re.compile(r'^\*\s+(.*)')
BLOCKQUOTE_PATTERN = re.compile(r'^>\s*(.*)')
CODE_BLOCK_PATTERN = re.compile(r'^```')
INLINE_CODE_PATTERN = re.compile(r'`(.*?)`')
BOLD_PATTERN = re.compile(r'\*\*(.*?)\*\*')
ITALIC_PATTERN = re.compile(r'\*(.*?)\*')
# Regular expression pattern for detecting Markdown hyperlinks
HYPERLINK_PATTERN = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
# HYPERLINK_PATTERN = re.compile(r'\[(.*?)\]\((.*?)\)')
# Image pattern looks for ![alt text](URL_or_path "title")
IMAGE_PATTERN = re.compile(r'!\[(.*?)\]\((.*?)\s*(?:"(.*?)")?\)')
# IMAGE_PATTERN = re.compile(r'!\[(.*?)\]\((.*?)\)')
TABLE_PATTERN = re.compile(r'^\|(.+)\|$')
    # Regular expression pattern for detecting Markdown table rows
TABLE_ROW_PATTERN = re.compile(r'\|(.+?)\|')
TABLE_DIVIDER_PATTERN = re.compile(r'\|\:?\-+\:?(\|\:?\-+\:?)*\|')

class MarkdownConverter:

    BULLET_PATTERN = re.compile(r'^\*\s+(.*)')

    def __init__(self, md_text):
        self.document = Document()
        self.md_text = md_text

    def convert(self):
        inside_code_block = False
        code_block_content = []

        for line in self.md_text.split('\n'):
            if self.start_or_end_code_block(line):
                inside_code_block = not inside_code_block
                if not inside_code_block:
                    self.handle_code_block(code_block_content)
                continue

            if inside_code_block:
                code_block_content.append(line)
                continue

            if self.handle_heading(line):
                continue
            if self.handle_bullet(line):
                continue
            # ... (call other handlers like handle_hyperlink, handle_blockquote etc.)
            
            self.handle_plain_text(line)

        return self.document

    def start_or_end_code_block(self, line):
        return line.strip().startswith('```')

    def add_code_style(self):
        # Define a character style named 'Code'
        char_style = self.document.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
        char_style.font.name = 'Consolas'
        char_style.font.size = Pt(10)
        char_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # Define a paragraph style named 'CodePara'
        para_style = self.document.styles.add_style('CodePara', WD_STYLE_TYPE.PARAGRAPH)
        para_style.font.name = 'Consolas'
        para_style.font.size = Pt(10)
        para_style.paragraph_format.line_spacing = 1.15  # Adjust line spacing as needed

    def handle_code_block(self, code_block_content):
        # Add the 'Code' style if it doesn't exist
        if 'Code' not in self.document.styles:
            self.add_code_style()

        # Create a new paragraph for the code block with the 'CodePara' style
        paragraph = self.document.add_paragraph(style='CodePara')

        # Add each line of the code block as a separate run
        for code_line in code_block_content:
            run = paragraph.add_run(code_line + '\n')
            run.style = 'Code'

        # Clear the content after it's added to the document
        code_block_content.clear()
    
    # def handle_code_block(self, code_block_content):
        # ... (implementation of code block handler)

    def handle_blockquote(self, lines, current_index):
        blockquote_content = []
        while current_index < len(lines) and lines[current_index].startswith('>'):
            # Strip the '> ' from the beginning of the line
            stripped_line = self.BLOCKQUOTE_PATTERN.match(lines[current_index]).group(1)
            blockquote_content.append(stripped_line)
            current_index += 1

        # Combine all the blockquote lines
        combined_blockquote = '\n'.join(blockquote_content)

        # Add the combined blockquote to the document with the desired style
        self.add_blockquote_to_document(combined_blockquote)

        # Return the updated index
        return current_index
    
    def add_blockquote_to_document(self, content):
        # Create a blockquote style if it does not exist
        # This should be a one-time setup in your style configuration code
        try:
            self.document.styles.add_style('BlockQuote', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:  # Style already exists
            pass

        # Apply the style to the paragraph
        paragraph = self.document.add_paragraph(content, style='BlockQuote')

        # Here you would format the paragraph to look like a blockquote
        # This may include setting the font, indentation, etc.        

    def handle_heading(self, line):
        match = HEADING_PATTERN.match(line)
        if match:
            level = len(match.group(1))  # Count the number of '#' characters
            text = match.group(2).strip()  # The actual text of the heading
            self.document.add_heading(text, level=level-1)  # Heading levels start with 0 in python-docx
            return True
        return False

    def handle_line(self, line):
        if self.inside_code_block:
            self.handle_code_block(line)
            return

        if CODE_BLOCK_PATTERN.match(line):
            self.inside_code_block = not self.inside_code_block
            if not self.inside_code_block:
                # Code block ending, add it to the document
                self.add_code_block_to_document()
            return

        # Check if the line is a heading
        if self.handle_heading(line):
            return

    # ... (handlers for other Markdown patterns)    

    def handle_bullet(self, line):
        match = self.BULLET_PATTERN.match(line)
        if match:
            bullet_text = match.group(1)
            self.add_bullet_to_document(bullet_text)
            return True
        return False

    def add_bullet_to_document(self, text):
        paragraph = self.document.add_paragraph(style='ListBullet')
        paragraph.add_run(text)

    #This method scans the text for markdown hyperlinks and adds them to the passed paragraph.
    def handle_hyperlink(self, paragraph, text): 
        cursor = 0
        for match in self.HYPERLINK_PATTERN.finditer(text):
            start, end = match.span()
            # Add text before the hyperlink
            if cursor < start:
                paragraph.add_run(text[cursor:start])

            # Extract hyperlink details
            link_text = match.group(1)
            link_url = match.group(2)

            # Add the hyperlink
            self.add_hyperlink(paragraph, link_url, link_text)

            # Move the cursor to the end of the hyperlink
            cursor = end
        
        # Add any remaining text after the last hyperlink
        if cursor < len(text):
            paragraph.add_run(text[cursor:])

    def handle_plain_text(self, line):
        # If the line is not empty, it's considered plain text
        if line.strip():
            self.add_plain_text_to_document(line.strip())
            return True
        return False

    def add_plain_text_to_document(self, text):
        paragraph = self.document.add_paragraph()
        paragraph.add_run(text)
        
    # ... (additional handlers for images, tables, blockquotes, etc.)
    
    def handle_list(self, lines, current_index):
        # Detect if the current line is a list item
        unordered_match = self.UNORDERED_LIST_PATTERN.match(lines[current_index])
        ordered_match = self.ORDERED_LIST_PATTERN.match(lines[current_index])

        # Process an unordered list
        if unordered_match:
            return self.handle_unordered_list(lines, current_index, unordered_match.group(1))

        # Process an ordered list
        if ordered_match:
            return self.handle_ordered_list(lines, current_index, ordered_match.group(1))

        # If it's not a list, return None
        return None

    def handle_unordered_list(self, lines, current_index, indent):
        # Start a list in the document
        self.document.add_paragraph(style='ListBullet')

        while current_index < len(lines):
            match = self.UNORDERED_LIST_PATTERN.match(lines[current_index])
            if match and match.group(1) == indent:
                # Add an item to the list
                paragraph = self.document.add_paragraph(match.group(3), style='ListBullet')
                current_index += 1
            else:
                # Not a list item, so end the list processing
                break

        # Return the updated index
        return current_index

    def handle_ordered_list(self, lines, current_index, indent):
        # Start a list in the document
        self.document.add_paragraph(style='ListNumber')

        while current_index < len(lines):
            match = self.ORDERED_LIST_PATTERN.match(lines[current_index])
            if match and match.group(1) == indent:
                # Add an item to the list
                paragraph = self.document.add_paragraph(match.group(2), style='ListNumber')
                current_index += 1
            else:
                # Not a list item, so end the list processing
                break

        # Return the updated index
        return current_index

    def handle_table(self, lines):
        # Determine where the table starts and ends
        table_start = None
        table_end = None
        for i, line in enumerate(lines):
            if self.TABLE_DIVIDER_PATTERN.match(line):
                if table_start is None:
                    table_start = i - 1
                table_end = i
            elif line.strip() == '' and table_start is not None:
                table_end = i
                break

        # If a table start was found but no end was detected, assume table goes to end of lines
        if table_start is not None and table_end is None:
            table_end = len(lines)

        # If no table was found, return False
        if table_start is None:
            return False

        # Extract header and rows
        header_row = lines[table_start]
        divider_row = lines[table_start + 1]
        body_rows = lines[table_start + 2:table_end]

        # Parse header row
        headers = self.TABLE_ROW_PATTERN.findall(header_row)
        num_columns = len(headers)

        # Create a table in the document
        table = self.document.add_table(rows=1, cols=num_columns)
        hdr_cells = table.rows[0].cells

        for i, header in enumerate(headers):
            hdr_cells[i].text = header.strip()

        # Add the rest of the rows
        for row in body_rows:
            row_data = self.TABLE_ROW_PATTERN.findall(row)
            if len(row_data) == num_columns:  # Ensure row has the correct number of columns
                row_cells = table.add_row().cells
                for i, cell in enumerate(row_data):
                    row_cells[i].text = cell.strip()

        # Return True and the index where the table ends
        return True, table_end

def handle_image(self, line):
        match = self.IMAGE_PATTERN.search(line)
        if match:
            alt_text = match.group(1)
            image_url_or_path = match.group(2)
            title = match.group(3) if len(match.groups()) > 2 else alt_text

            # You can adjust the width of the image by changing the value of Inches
            # For instance, to make the image width 4 inches you can call Inches(4)
            try:
                # Check if it's a URL or a file path
                if image_url_or_path.startswith('http'):  # Assume URL
                    response = urlopen(image_url_or_path)
                    image_stream = BytesIO(response.read())
                else:  # Assume file path
                    image_stream = image_url_or_path

                self.document.add_picture(image_stream, width=Inches(4))
                last_paragraph = self.document.paragraphs[-1]
                last_paragraph.alignment = 1  # Center the image

                if title:
                    self.document.add_paragraph(title).alignment = 1  # Center the title/caption

                return True
            except Exception as e:
                print(f"An error occurred while adding the image: {e}")
                # You might want to log the error or handle it according to your needs
                return False
        return False

    # A function that places a hyperlink within a paragraph object.
def add_hyperlink(self, paragraph, url, text):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, 'hyperlink', is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a new run object and add text
        new_run = paragraph.add_run()
        r = new_run._r
        r.append(hyperlink)

        # Create a new run object for the hyperlink text and apply style
        new_run = OxmlElement('w:r')
        r_style = OxmlElement('w:rStyle')
        r_style.set(qn('w:val'), 'Hyperlink')  # Set the style for hyperlinks
        new_run.append(r_style)
        new_run.text = text
        hyperlink.append(new_run)

# The function to convert a markdown file to a docx file
def convert_md_to_docx(md_filepath, docx_filepath):
    with open(md_filepath, 'r', encoding='utf-8') as file:
        md_text = file.read()

    converter = MarkdownConverter(md_text)
    doc = converter.convert()
    doc.save(docx_filepath)
    print(f"Converted {md_filepath} to {docx_filepath} successfully.")

# Example usage
if __name__ == "__main__":
    md_file = 'example.md'  # Your markdown file
    docx_file = 'example.docx'  # The Word file you want to create
    convert_md_to_docx(md_file, docx_file)
